from docx import Document
from copy import deepcopy
import sys,base64
from socket import inet_aton
import struct,os


def copy_table_after(ob_table, table):
    tbl, tb2 = ob_table._tbl, table._tbl
    new_tbl = deepcopy(tbl)
    tb2.addnext(new_tbl)


def copy_table_after3(document):
    p1, t1 = document.paragraphs[0]._p, document.tables[len(document.tables) - 4]._tbl
    new_tbl = deepcopy(p1)
    t1.addnext(new_tbl)


def copy_table_after2(pragraphs, table):
    p, tb = pragraphs._p, table._tbl
    new_tbl = deepcopy(tb)
    p.addnext(new_tbl)


def row_count(document):
    target_table = document.tables[len(document.tables) - 4]
    target_table.add_row()


def add_table(document):
    copy_table_after2(document.paragraphs[len(document.paragraphs) - 4], document.tables[len(document.tables) - 3])
    copy_table_after3(document)

def insert_data_main(document,ip,host,os):
    table=document.tables[len(document.tables)-4]
    table.cell(0,0).text=ip
    table.cell(1,0).text=host
    table.cell(3,0).text=ip
    table.cell(3,1).text=os


def insert_data_cvs(document,critical,high,med,low):
    table = document.tables[len(document.tables) - 4]
    print '-------->',critical,high,med,low
    table.cell(3,2).text=str(critical)
    table.cell(3,3).text=str(high)
    table.cell(3,4).text=str(med)
    table.cell(3,5).text=str(low)

def insert_data(document,port,cve,vuln,cvs):
    table=document.tables[len(document.tables)-4]
    print len(document.tables)
    row_s=table.rows
    table.cell(len(row_s)-1,3).text=vuln
    table.cell(len(row_s)-1,2).text=port
    table.cell(len(row_s)-1,1).text=cve
    table.cell(len(row_s)-1,0).text=cvs
    table.cell(len(row_s)-1,4).text=str(len(row_s)-2)

