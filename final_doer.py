import csv
import xmltodict,sys,sorter
from docx import Document
import json,import_to_word

file_name=sys.argv[1]
with open(file_name, 'rb') as f:
    reader = csv.reader(f)
    your_list = list(reader)
    your_list=your_list[::-1]
your_list.remove(your_list[0:1:][0])
#print your_list
document = Document('Nessus_table.docx')
huge= sorter.sorter(your_list)
def nuller(port, cve, vuln, cvs):
    out=[]
    if port=='':port='-'
    if cvs=='':cvs='-'
    if cve=='':cve='-'
    if vuln=='':vuln='-'
    out.append(cve)
    out.append(cvs)
    out.append(vuln)
    out.append(port)
    return out

def summer(per):
    med=0
    low=0
    critical=0
    high=0
    oza=[]
    for item in per:
        cvs = item[3]
        if cvs == 'Medium': med += 1
        if cvs == 'Low': low += 1
        if cvs == 'Critical': critical += 1
        if cvs == 'High': high += 1
    oza.append(low)
    oza.append(med)
    oza.append(high)
    oza.append(critical)
    return oza
def importer_to_docx(per):
        import_to_word.add_table(document)
        import_to_word.insert_data_main(document, per[0], per[1], per[2])
        per=per[3::]
        ret = per
        oza = summer(ret)
        #print oza

        try:
            import_to_word.insert_data_cvs(document, oza[3], oza[2], oza[1], oza[0])
        except Exception, e:
            print e

        import_to_word.copy_table_after(document.tables[len(document.tables) - 2],document.tables[len(document.tables) - 4])
        i=0
        for item in per:
            i+=1
            out=nuller(item[1],item[0],item[2],item[3])
            #print out
            import_to_word.row_count(document)
            import_to_word.insert_data(document, out[3], out[0], out[2], out[1])



out_file=file_name.replace('csv','docx')
i=0
def hig_to_low(a):
    ob=a[3::]
    def sortSecond(val):
        try:
            return float(val[4])
        except Exception,e:
            print e
    ob.sort(key=sortSecond)
    ob.reverse()
    out=[]
    out.append(a[0])
    out.append(a[1])
    out.append(a[2])
    for i in ob:
        out.append(i)
    print out
    return out
while huge!=[]:
    one= sorter.ye_cut(huge)
    out=hig_to_low(one)
    print out
    importer_to_docx(out)
    try:
        huge=huge.remove(out[0])
        document.save(out_file)
        exit(0)
    except:
        document.save(out_file)
        pass
    if huge==[]:
        break
    document.save(out_file)

